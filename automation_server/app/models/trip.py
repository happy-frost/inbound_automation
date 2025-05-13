import datetime
import copy

# import db from app
from app import db

from . import exceptions
from docx import Document


class Ticket(db.Model):
    __tablename__ = 'tickets'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    date = db.Column(db.Date, nullable=False)
    bought = db.Column(db.Boolean, nullable=False)
    price = db.Column(db.Float, nullable=True)
    trip_id = db.Column(db.Integer, db.ForeignKey('trips.id'), nullable=False)

    def __init__(self,name: str,date:str):
        self.name = name
        self.date = datetime.datetime.strptime(date,"%d %b %y")
        self.bought = False

class Trip(db.Model):
    __tablename__ = 'trips'

    id = db.Column(db.Integer, primary_key=True)
    guest_name = db.Column(db.String(100),nullable=False)
    start_date = db.Column(db.Date,nullable=False)
    end_date = db.Column(db.Date,nullable=False)
    adults = db.Column(db.Integer,nullable=False)
    children = db.Column(db.Integer,nullable=False)
    children_below_six = db.Column(db.Integer,nullable=False)
    attraction_tickets = db.relationship('Ticket', backref='trip', lazy=True, cascade='all, delete-orphan')
    whatsapp_group_id = db.Column(db.String(100),nullable=True)

    def __init__(self):
        self.guest_name = ""
        self.adults = 0
        self.children = 0
        self.children_below_six = 0
        self.attraction_tickets = []
    
    def populate_manually(self, guest_name, adults, children, children_below_six,attraction_tickets,start_date,end_date):
        self.guest_name = guest_name
        self.adults = adults
        self.children = children
        self.children_below_six = children_below_six
        self.start_date = start_date
        self.end_date = end_date

        for ticket in attraction_tickets:
            self.attraction_tickets.append(Ticket(ticket.name,ticket.date))

    def populate_with_docx(self,file_path):
        try:
            document = Document(file_path)
        except:
            raise exceptions.Document("there is an error opening the document")
        
        parsing_guest = False
        parsing_ticket = False
        parsing_dates = False

        for paragraph in document.paragraphs:
            if paragraph.text.startswith("Guest name"):
                guest = paragraph.text.split("\t")[1].split(" X ")
                self.guest_name = ' '.join(guest[0].split(" ")[1:]) # to remove the mr. or ms.
                guest_num = guest[1] # get the guest num from guest name
                parsing_guest = True
                continue

            if paragraph.text.startswith("Flight"):
                parsing_dates = True
                continue

            if parsing_guest:
                if paragraph.text.__contains__("+"):
                    guest_num += paragraph.text.strip("\t")
                else:
                    parsing_guest = False
            
            if parsing_dates:
                if paragraph.text == "":
                    parsing_dates = False
                    continue
                
                date = paragraph.text.split(" ")
                if date[2] == "ARR":
                    self.start_date = datetime.datetime.strptime(' '.join(date[:2] + [datetime.date.today().strftime("%y")]),"%d %b %y")
                    if self.start_date < datetime.datetime.now():
                        self.start_date = datetime.date(self.start_date.year + 1, self.start_date.month, self.start_date.day)
                if date[2] == "DEP":
                    self.end_date = datetime.datetime.strptime(' '.join(date[:2] + [datetime.date.today().strftime("%y")]),"%d %b %y")
                    if self.end_date < datetime.datetime.now():
                        self.end_date = datetime.date(self.end_date.year + 1, self.end_date.month, self.end_date.day)
                    
            if paragraph.text.startswith("Tickets"):
                guest_num_ticket = paragraph.text.split(" for ")[1] # get the guest num from tickets header
                parsing_ticket = True
                continue

            if parsing_ticket:
                if paragraph.text == "":
                    parsing_ticket = False
                    continue
                date, tickets_for_the_date = paragraph.text.strip(".").split(":")
                tickets_for_the_date = tickets_for_the_date.strip().split(", ")
                for item in tickets_for_the_date:
                    self.attraction_tickets.append(Ticket(item,date))
        
        if guest_num != guest_num_ticket:
            raise exceptions.Mismatch("there is a mismatch in the number of guest")

        guests = guest_num.split("+")
        for category in guests:
            if category.__contains__("adults"):
                self.adults = int(category.split(" ")[0])
            
            if category.__contains__("child"):
                self.children = int(category.strip().split(" ")[0])
                children_age_str = category.strip().split("(")[-1].strip("years)").split(" / ")
                children_age = list(map(lambda x: int(x.strip()), children_age_str))

                for age in children_age:
                    if age <= 6:
                        self.children_below_six += 1
    
    def add_whatsapp_group_id(self,group_id:str):
        self.whatsapp_group_id = group_id

    def tickets_to_buy(self):
        tickets_to_buy = copy.deepcopy(self.attraction_tickets)
        