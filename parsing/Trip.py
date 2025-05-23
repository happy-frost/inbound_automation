import datetime
import copy

# from flask_sqlalchemy import SQLAlchemy
from automation_server import db

from . import exceptions
from docx import Document

# db = SQLAlchemy()

class Ticket(db.Model):
    __tablename__ = 'tickets'
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100), nullable=False)
    date = db.Column(db.Date, nullable=False)
    bought = db.Column(db.Boolean, nullable=False)
    price = db.Column(db.Float, nullable=True)
    trip_id = db.Column(db.Integer, db.ForeignKey('trips.id'), nullable=False)

    def __init__(self,name: str,date:str):
        self.name = name
        self.date = datetime.datetime.strptime(date,"%d %b %y")
        self.bought = False

class Trip(db.Model):
    __tablename__ = 'trips'

    id = db.Column(db.Integer, primary_key=True)
    guest_name = db.Column(db.String(100),nullable=False)
    start_date = db.Column(db.Date,nullable=False)
    end_date = db.Column(db.Date,nullable=False)
    adults = db.Column(db.Integer,nullable=False)
    children = db.Column(db.Integer,nullable=False)
    children_below_six = db.Column(db.Integer,nullable=False)
    attraction_tickets = db.relationship('Ticket', backref='trip', lazy=True, cascade='all, delete-orphan')
    whatsapp_group_id = db.Column(db.String(100),nullable=True)

    def __init__(self,file_path):
        self.guest_name = ""
        self.start_date = None
        self.end_date = None
        self.adults = 0
        self.children = 0
        self.children_above_six = 0
        self.children_below_six = 0
        self.attraction_tickets = []
        
        try:
            document = Document(file_path)
        except:
            raise exceptions.Document("there is an error opening the document")
        
        parsing_guest = False
        parsing_ticket = False
        parsing_dates = False

        for paragraph in document.paragraphs:
            if paragraph.text.startswith("Guest name"):
                guest = paragraph.text.split("\t")[1].split(" X ")
                self.guest_name = guest[0]
                guest_num = guest[1] # get the guest num from guest name
                parsing_guest = True
                continue


            if parsing_guest:
                if paragraph.text.__contains__("+"):
                    guest_num += paragraph.text.strip("\t")
                else:
                    parsing_guest = False
            
            if paragraph.text.startswith("Flight"):
                parsing_dates = True
                continue

            if parsing_dates:
                if paragraph.text == "":
                    parsing_dates = False
                    continue
                date = paragraph.text.split(" ")
                if date[2] == "ARR":
                    self.start_date = datetime.datetime.strptime(''.join(date[:2]),"%d %b")
                if date[2] == "DEP":
                    self.end_date = datetime.datetime.strptime(''.join(date[:2]),"%d %b")
                    print(self.end_date)

            if paragraph.text.startswith("Tickets"):
                guest_num_ticket = paragraph.text.split(" for ")[1] # get the guest num from tickets header
                parsing_ticket = True
                continue

            if parsing_ticket:
                if paragraph.text == "":
                    parsing_ticket = False
                    continue
                date, tickets_for_the_date = paragraph.text.strip(".").split(":")
                tickets_for_the_date = tickets_for_the_date.strip().split(", ")
                for item in tickets_for_the_date:
                    self.attraction_tickets.append(Ticket(item,date))
        
        if guest_num != guest_num_ticket:
            raise exceptions.Mismatch("there is a mismatch in the number of guest")

        guests = guest_num.split("+")
        for category in guests:
            if category.__contains__("adults"):
                self.adults = int(category.split(" ")[0])
            
            if category.__contains__("child"):
                self.children = int(category.strip().split(" ")[0])
                children_age_str = category.strip().split("(")[-1].strip("years)").split(" / ")
                children_age = list(map(lambda x: int(x.strip()), children_age_str))

                for age in children_age:
                    if age <= 6:
                        self.children_below_six += 1
        
    def tickets_to_buy(self):
        tickets_to_buy = copy.deepcopy(self.attraction_tickets)
        